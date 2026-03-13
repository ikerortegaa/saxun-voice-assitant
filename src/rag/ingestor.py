"""
Pipeline de ingestión de documentos para el RAG de Saxun.
Soporta: PDF, DOCX, HTML, TXT/MD
Incluye: parsing, chunking, embeddings, deduplicación, versionado.
"""
import hashlib
import json
import mimetypes
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

import asyncpg
from loguru import logger

from src.config import get_settings
from src.models.rag_models import DocumentRecord
from src.rag.chunker import SemanticChunker, RawChunk
from src.rag.embeddings import EmbeddingService
from src.security.pii_redactor import get_redactor


class DocumentIngestor:
    """
    Ingesta documentos de rag-docs/ en la base de datos vectorial.
    Gestiona versionado, deduplicación y metadata de freshness.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".txt", ".md"}

    def __init__(self, db_pool: asyncpg.Pool):
        self._db = db_pool
        self._chunker = SemanticChunker()
        self._embedder = EmbeddingService()
        self._redactor = get_redactor()
        self._settings = get_settings()

    # ── Ingestión de un solo documento ────────────────────────────────────────

    async def ingest_file(
        self,
        file_path: str | Path,
        metadata: Optional[dict] = None,
    ) -> DocumentRecord:
        """
        Ingesta un documento completo.
        Si ya existe y no ha cambiado, lo omite.
        Si ha cambiado, marca la versión anterior como superseded.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Documento no encontrado: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Formato no soportado: {ext}")

        meta = metadata or {}
        file_hash = self._compute_hash(path)
        doc_id = self._generate_doc_id(path)

        # Verificar si ya existe y no ha cambiado
        existing = await self._get_existing_doc(doc_id)
        if existing and existing["file_hash"] == file_hash and existing["status"] == "active":
            logger.info(f"Documento sin cambios, omitiendo: {path.name}")
            return DocumentRecord(**existing)

        logger.info(f"Ingestando: {path.name}")

        # 1. Parsear documento
        raw_text, detected_lang = self._parse_document(path)
        if not raw_text.strip():
            raise ValueError(f"Documento vacío o no parseable: {path.name}")

        # 2. Detectar prompt injection en el documento
        if self._redactor.has_injection_attempt(raw_text):
            logger.warning(f"⚠️  Injection detectado en documento: {path.name} — RECHAZADO")
            raise SecurityError(f"Contenido sospechoso en {path.name}")

        # 3. Determinar tipo de documento para chunking
        doc_type = meta.get("doc_type", self._infer_doc_type(path, raw_text))
        language = meta.get("language", detected_lang)

        # 4. Chunking
        raw_chunks = self._chunker.chunk_document(raw_text, doc_type, language)
        logger.debug(f"  → {len(raw_chunks)} chunks generados")

        # 5. Embeddings en lote
        texts = [c.content for c in raw_chunks]
        embeddings = await self._embedder.embed_batch(texts)

        # 6. Crear registro de documento
        doc_record = DocumentRecord(
            doc_id=doc_id,
            file_path=str(path),
            file_hash=file_hash,
            title=meta.get("title", self._extract_title(raw_text, path.stem)),
            version=meta.get("version", "1.0"),
            status="active",
            language=language,
            sensitivity=meta.get("sensitivity", "public"),
            effective_date=meta.get("effective_date"),
            expiry_date=meta.get("expiry_date"),
            chunk_count=len(raw_chunks),
            metadata=meta,
        )

        # 7. Transacción: supersede versión anterior → insertar nueva
        async with self._db.acquire() as conn:
            async with conn.transaction():
                if existing:
                    await conn.execute(
                        "UPDATE document_registry SET status = 'superseded' WHERE doc_id = $1",
                        doc_id,
                    )
                    await conn.execute(
                        "UPDATE chunks SET status = 'superseded' WHERE doc_id = $1",
                        doc_id,
                    )

                # Insertar/actualizar registro de documento
                await self._upsert_document_record(conn, doc_record)

                # Insertar chunks con embeddings
                deduped = await self._deduplicate_chunks(conn, raw_chunks, doc_id)
                await self._insert_chunks(conn, deduped, embeddings, doc_record, raw_chunks)

        logger.success(f"✓ Ingestado: {path.name} → {len(raw_chunks)} chunks (doc_id: {doc_id})")
        return doc_record

    async def ingest_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
    ) -> list[DocumentRecord]:
        """Ingesta todos los documentos de un directorio."""
        base_path = Path(directory)
        if not base_path.exists():
            raise FileNotFoundError(f"Directorio no encontrado: {directory}")

        pattern = "**/*" if recursive else "*"
        files = [
            f for f in base_path.glob(pattern)
            if f.is_file() and f.suffix.lower() in self.SUPPORTED_EXTENSIONS
            and not f.name.startswith("_")
        ]

        logger.info(f"Ingestando directorio: {base_path} ({len(files)} archivos)")
        records = []
        errors = []

        for file_path in sorted(files):
            try:
                meta = self._load_metadata_file(file_path)
                record = await self.ingest_file(file_path, meta)
                records.append(record)
            except Exception as e:
                logger.error(f"Error ingestando {file_path.name}: {e}")
                errors.append((file_path, str(e)))

        logger.info(f"Ingestión completa: {len(records)} OK, {len(errors)} errores")
        return records

    # ── Parsing ───────────────────────────────────────────────────────────────

    def _parse_document(self, path: Path) -> tuple[str, str]:
        """Parsea el documento y devuelve (texto, idioma detectado)."""
        ext = path.suffix.lower()

        if ext == ".pdf":
            text = self._parse_pdf(path)
        elif ext == ".docx":
            text = self._parse_docx(path)
        elif ext in (".html", ".htm"):
            text = self._parse_html(path)
        else:  # .txt, .md
            text = path.read_text(encoding="utf-8", errors="replace")

        language = self._detect_language(text)
        return text, language

    def _parse_pdf(self, path: Path) -> str:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            page_text = page.get_text("text")
            if page_text.strip():
                pages.append(page_text)
        doc.close()
        return "\n\n".join(pages)

    def _parse_docx(self, path: Path) -> str:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Añadir indicador de encabezado
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    parts.append(f"{'#' * int(level)} {para.text}")
                else:
                    parts.append(para.text)
        # Incluir tablas
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n".join(rows))
        return "\n\n".join(parts)

    def _parse_html(self, path: Path) -> str:
        from bs4 import BeautifulSoup
        html = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")
        # Eliminar scripts y estilos
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    # ── Utilidades ────────────────────────────────────────────────────────────

    @staticmethod
    def _compute_hash(path: Path) -> str:
        sha256 = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    @staticmethod
    def _generate_doc_id(path: Path) -> str:
        """ID estable basado en la ruta del archivo (sin versión)."""
        normalized = path.stem.lower()
        normalized = re.sub(r'[^a-z0-9\-_]', '-', normalized)
        normalized = re.sub(r'-+', '-', normalized).strip('-')
        return normalized[:80]

    @staticmethod
    def _extract_title(text: str, fallback: str) -> str:
        lines = text.strip().split('\n')
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 5 and len(line) < 200:
                return re.sub(r'^#{1,4}\s+', '', line).strip()
        return fallback.replace('-', ' ').replace('_', ' ').title()

    @staticmethod
    def _infer_doc_type(path: Path, text: str) -> str:
        name = path.stem.lower()
        if "faq" in name or "preguntas" in name:
            return "faq"
        if "catalogo" in name or "productos" in name:
            return "catalog"
        if "politica" in name or "condicion" in name or "garantia" in name:
            return "policy"
        return "general"

    @staticmethod
    def _detect_language(text: str) -> str:
        try:
            from langdetect import detect
            sample = text[:2000]
            lang = detect(sample)
            return lang if lang in ("es", "ca", "en") else "es"
        except Exception:
            return "es"

    def _load_metadata_file(self, doc_path: Path) -> dict:
        """Busca un archivo .meta.json junto al documento."""
        meta_path = doc_path.with_suffix(".meta.json")
        if meta_path.exists():
            with open(meta_path) as f:
                return json.load(f)
        return {}

    async def _get_existing_doc(self, doc_id: str) -> Optional[dict]:
        async with self._db.acquire() as conn:
            row = await conn.fetchrow(
                "SELECT * FROM document_registry WHERE doc_id = $1", doc_id
            )
            if row is None:
                return None
            result = dict(row)
            if isinstance(result.get("metadata"), str):
                result["metadata"] = json.loads(result["metadata"])
            return result

    async def _upsert_document_record(self, conn, record: DocumentRecord) -> None:
        await conn.execute("""
            INSERT INTO document_registry
                (doc_id, file_path, file_hash, title, version, status, language,
                 sensitivity, effective_date, expiry_date, chunk_count, ingested_at, metadata)
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12,$13)
            ON CONFLICT (doc_id) DO UPDATE SET
                file_path=EXCLUDED.file_path, file_hash=EXCLUDED.file_hash,
                title=EXCLUDED.title, version=EXCLUDED.version,
                status=EXCLUDED.status, language=EXCLUDED.language,
                sensitivity=EXCLUDED.sensitivity, effective_date=EXCLUDED.effective_date,
                expiry_date=EXCLUDED.expiry_date, chunk_count=EXCLUDED.chunk_count,
                ingested_at=EXCLUDED.ingested_at, metadata=EXCLUDED.metadata
        """,
            record.doc_id, record.file_path, record.file_hash, record.title,
            record.version, record.status, record.language, record.sensitivity,
            record.effective_date, record.expiry_date, record.chunk_count,
            record.ingested_at, json.dumps(record.metadata),
        )

    async def _deduplicate_chunks(
        self, conn, chunks: list[RawChunk], doc_id: str
    ) -> list[RawChunk]:
        """Elimina chunks casi idénticos a contenido ya existente en la DB."""
        # Deduplicación simple basada en hash de contenido
        seen_hashes = set()
        unique = []
        for chunk in chunks:
            content_hash = hashlib.md5(chunk.content.encode()).hexdigest()
            if content_hash not in seen_hashes:
                seen_hashes.add(content_hash)
                unique.append(chunk)
        removed = len(chunks) - len(unique)
        if removed:
            logger.debug(f"  Deduplicación: {removed} chunks duplicados eliminados")
        return unique

    async def _insert_chunks(
        self,
        conn,
        chunks: list[RawChunk],
        embeddings: list[list[float]],
        doc_record: DocumentRecord,
        original_chunks: list[RawChunk],
    ) -> None:
        """Inserta chunks con embeddings en la base de datos."""
        import numpy as np
        for i, chunk in enumerate(chunks):
            chunk_id = chunk.generate_id(doc_record.doc_id)
            embedding = embeddings[i] if i < len(embeddings) else embeddings[-1]
            # pgvector.asyncpg.register_vector espera numpy ndarray
            embedding_vec = np.array(embedding, dtype=np.float32)
            metadata = {**chunk.metadata, **doc_record.metadata}

            await conn.execute("""
                INSERT INTO chunks
                    (chunk_id, doc_id, content, section, language, sensitivity,
                     status, chunk_index, embedding, metadata, created_at)
                VALUES ($1,$2,$3,$4,$5,$6,'active',$7,$8,$9,$10)
                ON CONFLICT (chunk_id) DO UPDATE SET
                    content=EXCLUDED.content, embedding=EXCLUDED.embedding,
                    status='active', metadata=EXCLUDED.metadata
            """,
                chunk_id, doc_record.doc_id, chunk.content, chunk.section,
                chunk.language, doc_record.sensitivity, chunk.chunk_index,
                embedding_vec, json.dumps(metadata), datetime.utcnow(),
            )


class SecurityError(Exception):
    pass
